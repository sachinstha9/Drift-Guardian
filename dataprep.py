"""
Document ingestion for DriftGuardian.

Supports two modes:
 1) Local filesystem (legacy): load policy/SOP markdown files from data/
 2) Uploaded files (new): accept PDF, DOCX, TXT, MD via multipart upload and
    extract text in-process.

For an OPEA-aligned deployment, the upload path can be redirected to the
OPEA dataprep microservice by setting OPEA_DATAPREP_URL — see ingest_via_opea().
"""
import io
import os
import uuid
import logging
from pathlib import Path
from typing import Optional

import httpx

logger = logging.getLogger(__name__)

DATA_DIR = Path(__file__).parent.parent / "data"
POLICY_DIR = DATA_DIR / "policy_hierarchy"
SOP_DIR = DATA_DIR / "sop_drafts"

# Optional OPEA dataprep microservice. If set, uploads are forwarded there.
# Example: http://dataprep:6007/v1/dataprep/ingest
OPEA_DATAPREP_URL = os.environ.get("OPEA_DATAPREP_URL", "").strip()

# In-memory document store keyed by doc_id. For production this would be Redis
# or a vector DB; for the demo it's fine.
_DOC_STORE: dict[str, dict] = {}


# ------------------------------------------------------------------ #
# Filesystem helpers (legacy demo path)
# ------------------------------------------------------------------ #
def load_text(filepath: Path) -> str:
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def load_global_baseline() -> str:
    return load_text(POLICY_DIR / "global_baseline_kyc_policy.md")


def load_regional_override(region: str) -> Optional[str]:
    region_map = {
        "APAC": "apac_override_policy.md",
        "EU": "eu_override_policy.md",
    }
    filename = region_map.get(region.upper())
    if not filename:
        return None
    path = POLICY_DIR / filename
    return load_text(path) if path.exists() else None


def load_sop(filename: str) -> str:
    path = SOP_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"SOP file not found: {filename}")
    return load_text(path)


def list_sop_files() -> list[str]:
    if not SOP_DIR.exists():
        return []
    return [f.name for f in SOP_DIR.glob("*.md")]


# ------------------------------------------------------------------ #
# Upload / extraction
# ------------------------------------------------------------------ #
def _extract_pdf(raw: bytes) -> str:
    """Extract text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf is required for PDF uploads. Install with `pip install pypdf`."
        ) from e

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 - per-page failure shouldn't abort
            logger.warning("PDF page extraction failed: %s", e)
    return "\n".join(parts).strip()


def _extract_docx(raw: bytes) -> str:
    """Extract text from a .docx using python-docx."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX uploads. Install with "
            "`pip install python-docx`."
        ) from e

    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text]
    # Include table cell text too — policy docs love tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def _extract_plain(raw: bytes) -> str:
    """Decode plaintext / markdown / html-as-text."""
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace").strip()


def extract_text_from_upload(filename: str, content_type: str, raw: bytes) -> str:
    """
    Route an uploaded file to the correct extractor based on extension and
    content-type. Returns plain text. Raises ValueError for unsupported types.
    """
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        return _extract_pdf(raw)
    if name.endswith(".docx") or "wordprocessingml" in ctype:
        return _extract_docx(raw)
    if name.endswith((".txt", ".md", ".markdown", ".html", ".htm")) \
            or ctype.startswith("text/"):
        return _extract_plain(raw)

    # Last-ditch: try decoding as text. Many policy docs are sent as
    # application/octet-stream but are actually UTF-8.
    try:
        return _extract_plain(raw)
    except Exception as e:  # noqa: BLE001
        raise ValueError(
            f"Unsupported file type: filename={filename!r} content_type={content_type!r}"
        ) from e


async def ingest_via_opea(
    filename: str, content_type: str, raw: bytes
) -> Optional[str]:
    """
    Forward an upload to the OPEA dataprep microservice if configured.
    Returns extracted text on success, None if OPEA dataprep isn't configured.
    Falls through to local extraction on error so the demo never hard-fails.
    """
    if not OPEA_DATAPREP_URL:
        return None

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            files = {"files": (filename, raw, content_type or "application/octet-stream")}
            resp = await client.post(OPEA_DATAPREP_URL, files=files)
            resp.raise_for_status()
            # OPEA dataprep responses vary by backend; try a few common shapes.
            data = resp.json()
            if isinstance(data, dict):
                for key in ("text", "content", "extracted_text", "data"):
                    if key in data and isinstance(data[key], str):
                        return data[key]
            logger.warning("OPEA dataprep returned unexpected shape: %s", type(data))
            return None
    except Exception as e:  # noqa: BLE001
        logger.warning(
            "OPEA dataprep call failed (%s); falling back to local extraction.", e
        )
        return None


async def store_upload(filename: str, content_type: str, raw: bytes) -> dict:
    """
    Ingest an uploaded file and store the extracted text under a doc_id.
    Tries OPEA dataprep first, falls back to local extraction.
    """
    text = await ingest_via_opea(filename, content_type, raw)
    if not text:
        text = extract_text_from_upload(filename, content_type, raw)

    if not text or not text.strip():
        raise ValueError(f"No text could be extracted from {filename!r}")

    doc_id = uuid.uuid4().hex[:12]
    _DOC_STORE[doc_id] = {
        "filename": filename,
        "content_type": content_type,
        "text": text,
    }
    return {
        "doc_id": doc_id,
        "filename": filename,
        "content_type": content_type,
        "chars": len(text),
        "preview": text[:300],
    }


def get_uploaded_text(doc_id: str) -> Optional[str]:
    entry = _DOC_STORE.get(doc_id)
    return entry["text"] if entry else None